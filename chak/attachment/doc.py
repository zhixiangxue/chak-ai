"""
Word document attachment with default reader.
"""
from typing import Optional, Callable, Union, Awaitable
import os
import tempfile
from .base import Attachment, MimeType, detect_content_type, AttachmentContent, ContentType


def default_docx_reader(content: str) -> AttachmentContent:
    """
    Default synchronous DOCX reader.
    
    Args:
        content: Local file path or URL to DOCX
        
    Returns:
        AttachmentContent with extracted text and metadata
    """
    try:
        from docx import Document
        import urllib.request
        
        content_type = detect_content_type(content)
        
        # Handle different content types
        if content_type == ContentType.URL:
            # Download URL to temp file
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                urllib.request.urlretrieve(content, tmp.name)
                file_path = tmp.name
        
        elif content_type == ContentType.LOCAL_PATH:
            # Local path: check if file exists
            if not os.path.exists(content):
                return AttachmentContent(content="", meta={"error": f"File not found: {content}"})
            file_path = content
        
        elif content_type == ContentType.TEXT:
            # Plain text: not a valid file path
            return AttachmentContent(content="", meta={"error": "Invalid content: expected file path or URL, got plain text"})
        
        else:  # ContentType.DATA_URI
            # Data URI: not supported for DOCX
            return AttachmentContent(content="", meta={"error": "Data URI not supported for DOCX files"})
        
        # Now file_path is guaranteed to be a valid local file
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        meta = {
            "paragraphs": len(doc.paragraphs),
        }
        
        # Clean up temp file if downloaded from URL
        if content_type == ContentType.URL:
            try:
                os.unlink(file_path)
            except Exception:
                pass
        
        return AttachmentContent(content=text, meta=meta)
    
    except Exception as e:
        return AttachmentContent(content="", meta={"error": str(e)})


async def default_docx_reader_async(content: str) -> AttachmentContent:
    """
    Default asynchronous DOCX reader.
    
    Args:
        content: Local file path or URL to DOCX
        
    Returns:
        AttachmentContent with extracted text and metadata
    """
    try:
        from docx import Document
        import aiohttp
        import aiofiles
        
        content_type = detect_content_type(content)
        
        # Handle different content types
        if content_type == ContentType.URL:
            # Download URL to temp file asynchronously
            async with aiohttp.ClientSession() as session:
                async with session.get(content) as response:
                    response.raise_for_status()
                    data = await response.read()
                    
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                        tmp.write(data)
                        file_path = tmp.name
        
        elif content_type == ContentType.LOCAL_PATH:
            # Local path: check if file exists
            if not os.path.exists(content):
                return AttachmentContent(content="", meta={"error": f"File not found: {content}"})
            file_path = content
        
        elif content_type == ContentType.TEXT:
            # Plain text: not a valid file path
            return AttachmentContent(content="", meta={"error": "Invalid content: expected file path or URL, got plain text"})
        
        else:  # ContentType.DATA_URI
            # Data URI: not supported for DOCX
            return AttachmentContent(content="", meta={"error": "Data URI not supported for DOCX files"})
        
        # Now file_path is guaranteed to be a valid local file
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        meta = {
            "paragraphs": len(doc.paragraphs),
        }
        
        # Clean up temp file if downloaded from URL
        if content_type == ContentType.URL:
            try:
                os.unlink(file_path)
            except Exception:
                pass
        
        return AttachmentContent(content=text, meta=meta)
    
    except Exception as e:
        return AttachmentContent(content="", meta={"error": str(e)})


class DOC(Attachment):
    """Microsoft Word document attachment (DOC/DOCX)"""
    
    def __init__(
        self, 
        source: str,
        reader: Optional[Callable[[str], Union[AttachmentContent, Awaitable[AttachmentContent]]]] = None
    ):
        """
        Create a Word document attachment.
        
        Args:
            source: Local file path or URL to DOC/DOCX file
            reader: Optional custom reader function (defaults to default_docx_reader)
        
        Examples:
            >>> DOC("document.docx")
            >>> DOC("https://example.com/file.docx")
        """
        # Auto-detect DOCX vs DOC from extension
        mime_type = MimeType.DOCX if source.lower().endswith('.docx') else MimeType.DOC
        reader = reader or default_docx_reader_async
        super().__init__(source, mime_type, reader)
